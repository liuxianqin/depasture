import os
import shutil
import zipfile
from docx import Document
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import docx

def copyTable(table, document):
    # 按照单元格复制表格
    R = len(table.rows)
    C = len(table.columns)
    new_table = document.add_table(rows=R, cols=C, style='Table Grid')
    for row_num in range( len(table.rows) ):
        for col_num in range( len(table.columns)):
            new_table.cell(row_num,col_num).text = table.cell(row_num,col_num).text

def iter_block_items(parent):
    """
    判断word中的是段落还是表格
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def getTablePos(document):

    # 获得段落和表格的位置关系  假设表格前面都存在一个段落
    # 返回 表格之前一个段落在 paragraphs中的段落号
    abs_paragraph_lst = []
    abs_table_lst = []
    for index, block in enumerate(iter_block_items(document)):
        # print(index,block)
        if "docx.text.paragraph.Paragraph object"  in str(block):
            abs_paragraph_lst.append(index)
        if "docx.table.Table object" in str(block):
            abs_table_lst.append(index)
    # print(abs_paragraph_lst)
    # print(abs_table_lst)
    result_lst = [    abs_paragraph_lst.index( i-1  )  for i in abs_table_lst         ]
    # print(result_lst)

    paras = []
    paras_nums = []
    for index in range(len(document.paragraphs)):
        if index in result_lst:
            # print("前一个   发现相同", document.paragraphs[index].text)
            paras.append(document.paragraphs[index])
            paras_nums.append(index)
    # return [    abs_paragraph_lst.index( i-1  )  for i in abs_table_lst         ]
    return paras,paras_nums

def word2pic(path, zip_path, tmp_path, store_path):
    """
    对被合并的docx做此操作
    :param path:源文件
    :param zip_path:docx重命名为zip
    :param tmp_path:中转图片文件夹
    :param store_path:最后保存结果的文件夹（需要手动创建）
    :return:
    """
    # 将docx文件重命名为zip文件
    os.rename(path, zip_path)
    # 进行解压
    f = zipfile.ZipFile(zip_path, 'r')
    # 将图片提取并保存
    for file in f.namelist():
        f.extract(file, tmp_path)
    # 释放该zip文件
    f.close()

    # 将docx文件从zip还原为docx
    os.rename(zip_path, path)
    # 得到缓存文件夹中图片列表
    pic = os.listdir(os.path.join(tmp_path, 'word/media'))

    # 将图片复制到最终的文件夹中
    sortoffset_num = 100000
    for i in pic:
        # 根据word的路径生成图片的名称
        sortoffset_num += 1
        new_name = path.replace('\\', '_')
        new_name = new_name.replace(':', '') + '_' + i
        new_name = str(sortoffset_num) + '_' + new_name
        # print(new_name)
        shutil.copy(os.path.join(tmp_path + '/word/media', i), os.path.join(store_path, new_name))

    # 删除缓冲文件夹中的文件，用以存储下一次的文件
    for i in os.listdir(tmp_path):
        # 如果是文件夹则删除
        if os.path.isdir(os.path.join(tmp_path, i)):
            shutil.rmtree(os.path.join(tmp_path, i))


def belong(document, HeadingText):
    """
    找到属于这个标题下的所有段落
    :param paragraphs:
    :param HeadingText:
    :return:
    """
    paragraphs = document.paragraphs
    para_list = []
    para_Indexes = []
    father_index = -1
    father_stylename = ""
    # 寻找父结点
    for index in range(len(paragraphs)):
        if paragraphs[index].text == HeadingText  and paragraphs[index].style.name == "Heading 1":
            father_index = index
            father_stylename = paragraphs[index].style.name
    if father_index >= 0:               # =0 是因为第一段的大标题序号是0
        for index in range(father_index + 1, len(paragraphs) - 1):
            # 知道遇见下一个同级别的段落终止，中间的是所有段落
            if paragraphs[index].style.name == father_stylename:
                break
            para_list.append(paragraphs[index])
            # para_list.append(paragraphs[index].text.replace(u'\u2002', u' '))  # 中文空格
    for index in range(len(paragraphs)):
        if paragraphs[index] in para_list:
            # print("相同")
            para_Indexes.append(index)
    PARA, PARA_index = getTablePos(document)
    # print(PARA)
    # print('PARA-index',PARA_index)
    # 再返回一个 index in paragraphs
    return para_list, para_Indexes, PARA, PARA_index


def getHeading1(paragraphs, index):
    # 得到大标题 返回大标题的列表
    structSmg = []
    for p in paragraphs:
        structSmg.append(p.style.name)
    for i in range(index, -1, -1):
        if structSmg[i] == "Heading 1":
            return i
    else:
        print("None")
        return None


def mergeDocx(filename1, filename2):
    # 按照顺序 把1、2打印到3

    # 先把docx1的图片保存
    word2pic(filename1, r'log.zip', r'./document1', r'./img1')

    # 先把docx2的图片保存
    word2pic(filename2, r'log.zip', r'./document2', r'./img2')

    document1 = docx.Document(filename1)
    document2 = docx.Document(filename2)
    document3 = docx.Document()


    # 读取表格
    tables1 = document1.tables
    tables2 = document2.tables
    # tables1_parasLst = getTablePos(document1)
    # tables2_parasLst = getTablePos(document2)

    # 找出docx1的图片所在的paragraph
    img_index_lst1 = []
    new_img_index_lst1 = []
    for i in range(len(document1.paragraphs)):
        if 'graphicData' in document1.paragraphs[i]._p.xml:
            img_index_lst1.append(i)
            # print('图片所在para:', '0' + document1.paragraphs[i].text + '$', document1.paragraphs[i].style.name)
    # print('图片总数',len(img_index_lst1), img_index_lst1)
    # 找出docx2的图片所在的paragraph
    img_index_lst2 = []
    new_img_index_lst2 = []
    for i in range(len(document2.paragraphs)):
        if 'graphicData' in document2.paragraphs[i]._p.xml:
            img_index_lst2.append(i)
            # print('图片所在para:', '0' + document2.paragraphs[i].text + '$', document2.paragraphs[i].style.name)
    # print('图片总数',len(img_index_lst2), img_index_lst2)


    # 构建按大标题排列的列表
    Heading1_list = []
    for para in document1.paragraphs:
        if para.style.name == "Heading 1":
            Heading1_list.append(para.text)
    # print('大标题',Heading1_list)


    # 按顺序 依次添加段落
    for Heading1 in Heading1_list:
        document3.add_paragraph(Heading1, "Heading 1")
        # 1 2 的正文       第三个返回值每次都改变  第四个是位置每次没有改变
        para_of_1 , paraIndex_of_1, PARA1, PARA_Index1 = belong(document1, Heading1)
        print(paraIndex_of_1,PARA_Index1)
        para_of_2 ,paraIndex_of_2, PARA2, PARA_Index2 = belong(document2, Heading1)
        print(paraIndex_of_2, PARA_Index2)
        print("::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::")


        # document3.add_paragraph("1属于本段的正文","Normal")

        for para, paraIndex_of_1 in zip(para_of_1, paraIndex_of_1):
            Down_Flag = False        # 是否在表格哪里已经添加过此段

            # 检查是否应该包含表格
            print(paraIndex_of_1,para.text)
            if paraIndex_of_1 in PARA_Index1:
                the_which = PARA_Index1.index(paraIndex_of_1)
                # print('the_which:', the_which)
                document3.add_paragraph(para.text, para.style.name)
                table = tables1[the_which]
                copyTable(table, document3)
                Down_Flag = True


            # 偏移计算图片在新文档中应该插入的位置
            if 'graphicData' in para._p.xml:
                # print("这里有docx1的一张照片 ——————————————————————————————————————————————————————————")
                document3.add_paragraph(para.text, para.style.name)
                new_img_index_lst1.append(len(document3.paragraphs)-1)
            else:
                if Down_Flag == False:
                    document3.add_paragraph(para.text, para.style.name)



        for para, paraIndex_of_2 in zip(para_of_2,paraIndex_of_2):

            Down_Flag = False  # 是否在表格哪里已经添加过此段
            # 检查是否应该包含表格
            print(paraIndex_of_2, para.text)
            if paraIndex_of_2 in PARA_Index2:
                the_which = PARA_Index2.index(paraIndex_of_2)
                # print('the_which:', the_which)
                document3.add_paragraph(para.text, para.style.name)
                table = tables2[the_which]
                copyTable(table, document3)
                Down_Flag = True

            if 'graphicData' in para._p.xml:
                # print("这里有docx2的一张照片 ——————————————————————————————————————————————————————————")
                document3.add_paragraph(para.text, para.style.name)
                new_img_index_lst2.append(len(document3.paragraphs)-1)
            else:
                if Down_Flag == False:
                    document3.add_paragraph(para.text,para.style.name)
        # document3.add_paragraph("2属于本段的正文","Normal")

    img_lst1 = os.listdir('./img1')
    # print(img_lst1)
    img_lst2 = os.listdir('./img2')
    # print(img_lst2)

    for i in range(len(img_lst1)):
        document3.paragraphs[new_img_index_lst1[i]].add_run().add_picture("img1/" + img_lst1[i])
    for i in range(len(img_lst2)):
        document3.paragraphs[new_img_index_lst2[i]].add_run().add_picture("img2/" + img_lst2[i])


    document3.save("合并后的文档.docx")
    return document3


if __name__ == "__main__":
    # filename1 = "关于增加移动Pad进件功能的需求v1.8.docx"
    # filename2 = "关于手机银行手机号转账的优化需求V1.0.docx"

    filename1 = "手机号转账.docx"
    filename2 = "移动Pad.docx"

    mergeDocx(filename1, filename2)

